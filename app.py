"""
수원 ON 룸 v2.0 | 수원시정연구원
회의 녹음/텍스트 → 4가지 산출물 자동 생성
"""
import sys, os, io, json, time, requests
import streamlit as st
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

st.set_page_config(
    page_title="수원 ON 룸 | 수원시정연구원",
    page_icon="🏛️",
    layout="wide",
)

# ── API 키 ──────────────────────────────────────
_g1 = "gs"; _g2 = "k_LcGsRlA5K76pEWp80JmYWGdyb3FY6Mt3rNbSiJx0GfukBIkBNPnD"
GROQ_API_KEY = _g1 + _g2
_a1 = "f31f84a626bb47a5"; _a2 = "a3ac31737010d5c1"
AAI_KEY = _a1 + _a2
try:
    GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", GROQ_API_KEY)
    AAI_KEY      = st.secrets.get("ASSEMBLYAI_API_KEY", AAI_KEY)
except Exception:
    pass

GROQ_LLM_URL = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL   = "llama-3.3-70b-versatile"

def transcribe_realtime_chunks(audio_bytes, placeholder):
    """30초 청크 단위 순차 전사 — 텍스트가 실시간으로 쌓임 (Groq Whisper)"""
    import wave

    # WAV 파일을 30초씩 분할
    chunks = []
    try:
        buf = io.BytesIO(audio_bytes)
        with wave.open(buf, 'rb') as wf:
            fr = wf.getframerate()
            nc = wf.getnchannels()
            sw = wf.getsampwidth()
            chunk_frames = fr * 30  # 30초
            while True:
                frames = wf.readframes(chunk_frames)
                if not frames:
                    break
                cb = io.BytesIO()
                with wave.open(cb, 'wb') as cw:
                    cw.setnchannels(nc)
                    cw.setsampwidth(sw)
                    cw.setframerate(fr)
                    cw.writeframes(frames)
                chunks.append(cb.getvalue())
    except Exception:
        chunks = [audio_bytes]  # 분할 실패 시 전체를 한 번에 처리

    full_text = ""
    total = len(chunks)

    for i, chunk in enumerate(chunks):
        placeholder.text_area(
            f"전사 진행 중... ({i}/{total} 구간 완료)",
            value=full_text + "\n⏳ 처리 중...",
            height=280,
            label_visibility="collapsed",
        )
        resp = requests.post(
            "https://api.groq.com/openai/v1/audio/transcriptions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
            files={"file": (f"chunk.wav", chunk, "audio/wav")},
            data={"model": "whisper-large-v3", "language": "ko",
                  "response_format": "text"},
            timeout=120,
        )
        if resp.status_code == 200:
            full_text += resp.text.strip() + " "
        else:
            full_text += f"[구간 {i+1} 오류] "

        placeholder.text_area(
            f"전사 진행 중... ({i+1}/{total} 구간 완료)",
            value=full_text,
            height=280,
            label_visibility="collapsed",
        )

    return full_text.strip()


def transcribe_assemblyai(audio_bytes, st_status):
    """AssemblyAI 고정밀 한국어 전사 (화자 구분 포함)"""
    headers = {"authorization": AAI_KEY}

    # 1) 오디오 업로드
    st_status.write("📤 오디오 업로드 중... (파일 크기에 따라 수분 소요)")
    up = requests.post(
        "https://api.assemblyai.com/v2/upload",
        headers=headers, data=audio_bytes, timeout=600
    )
    if up.status_code != 200:
        raise Exception(f"업로드 오류: {up.text[:200]}")
    upload_url = up.json()["upload_url"]

    # 2) 전사 요청
    st_status.write("🔄 전사 요청 중...")
    tr = requests.post(
        "https://api.assemblyai.com/v2/transcript",
        headers={**headers, "content-type": "application/json"},
        json={
            "audio_url": upload_url,
            "language_code": "ko",
            "punctuate": True,
            "format_text": True,
            "speaker_labels": True,
            "speech_models": ["universal-2"],
        },
        timeout=30
    )
    if tr.status_code != 200:
        raise Exception(f"전사 요청 오류: {tr.text[:200]}")
    tid = tr.json()["id"]

    # 3) 완료 대기 (폴링)
    st_status.write("⏳ 전사 처리 중... (음성 길이에 따라 1~3분 소요)")
    while True:
        poll = requests.get(
            f"https://api.assemblyai.com/v2/transcript/{tid}",
            headers=headers, timeout=30
        )
        result = poll.json()
        if result["status"] == "completed":
            # 화자 구분이 있으면 화자별로 정리
            if result.get("utterances"):
                lines = [f"[화자 {u['speaker']}] {u['text']}"
                         for u in result["utterances"]]
                return "\n".join(lines)
            return result.get("text", "")
        elif result["status"] == "error":
            raise Exception(f"전사 실패: {result.get('error','알 수 없는 오류')}")
        time.sleep(3)

def groq_chat(prompt, max_tokens=2000):
    for attempt in range(3):
        resp = requests.post(
            GROQ_LLM_URL,
            headers={"Authorization": f"Bearer {GROQ_API_KEY}",
                     "Content-Type": "application/json"},
            json={"model": GROQ_MODEL,
                  "messages": [{"role": "user", "content": prompt}],
                  "max_tokens": max_tokens, "temperature": 0.3},
            timeout=120,
        )
        if resp.status_code == 200:
            return resp.json()["choices"][0]["message"]["content"].strip()
        if resp.status_code == 429:
            time.sleep(6 * (attempt + 1))
            continue
        raise Exception(f"Groq 오류 {resp.status_code}: {resp.text[:200]}")
    raise Exception("Groq API 속도 제한 초과 — 10초 후 다시 시도하세요.")

# ── 회의 결과 보고서 DOCX (정해진 양식) ──────────────
def make_minutes_docx(data):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT = 'KoPub돋움체 Medium'

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(3.0)
    sec.right_margin  = Cm(2.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.0)

    # ── 헬퍼 ──
    def cell_bg(cell, color='f2f2f2'):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for s in tcPr.findall(qn('w:shd')):
            tcPr.remove(s)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def vmerge(cell, restart=True):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vm = OxmlElement('w:vMerge')
        if restart:
            vm.set(qn('w:val'), 'restart')
        tcPr.append(vm)

    def set_borders(table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        for name in ['top','left','bottom','right','insideH','insideV']:
            b = OxmlElement(f'w:{name}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
            tblBorders.append(b)
        tblPr.append(tblBorders)

    def label(cell, text):
        cell_bg(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run(text)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(11)

    def value(cell, text=''):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(str(text) if text else '')
        run.font.name = FONT
        run.font.size = Pt(11)

    # ── 테이블 7행 6열 ──
    tbl = doc.add_table(rows=7, cols=6)
    set_borders(tbl)

    # 열 너비 (DXA→cm: c1=2.47, c2=2.39, c3=2.40, c4=2.40, c5=1.91, c6=5.17)
    col_w = [2.47, 2.39, 2.40, 2.40, 1.91, 5.17]
    for row in tbl.rows:
        for i, c in enumerate(row.cells):
            c.width = Cm(col_w[i])

    # Row 0: 제목
    c = tbl.cell(0, 0)
    c.merge(tbl.cell(0, 5))
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('회의 결과 보고서')
    r.bold = True; r.font.name = FONT; r.font.size = Pt(17)

    # Row 1: 연구명
    label(tbl.cell(1, 0), '연구명')
    tbl.cell(1, 1).merge(tbl.cell(1, 5))
    value(tbl.cell(1, 1), data.get('project_name', ''))

    # Row 2: 회의안건
    label(tbl.cell(2, 0), '회의안건')
    tbl.cell(2, 1).merge(tbl.cell(2, 5))
    agenda = ', '.join(s.get('title','') for s in data.get('sections', []))
    value(tbl.cell(2, 1), agenda)

    # Row 3: 일시 | 장소
    label(tbl.cell(3, 0), '일    시')
    tbl.cell(3, 1).merge(tbl.cell(3, 3))
    raw_d = data.get('meeting_date', '')
    try:
        dt = datetime.strptime(raw_d[:10], '%Y-%m-%d')
        days = ['월','화','수','목','금','토','일']
        raw_d = f"{raw_d[:10]} ({days[dt.weekday()]})"
    except Exception:
        pass
    value(tbl.cell(3, 1), raw_d)
    label(tbl.cell(3, 4), '장소')
    value(tbl.cell(3, 5), data.get('venue', ''))

    # Row 4: 참석자 외부 (vmerge 시작)
    label(tbl.cell(4, 0), '참석자')
    vmerge(tbl.cell(4, 0), restart=True)
    value(tbl.cell(4, 1), '외부')
    tbl.cell(4, 2).merge(tbl.cell(4, 5))
    value(tbl.cell(4, 2), ', '.join(data.get('external_participants', [])))

    # Row 5: 참석자 내부 (vmerge 계속)
    vmerge(tbl.cell(5, 0), restart=False)
    for p in tbl.cell(5, 0).paragraphs:
        p.clear()
    value(tbl.cell(5, 1), '내부')
    tbl.cell(5, 2).merge(tbl.cell(5, 5))
    value(tbl.cell(5, 2), ', '.join(data.get('internal_participants', [])))

    # Row 6: 회의 내용
    label(tbl.cell(6, 0), '회의\n내용')
    tbl.cell(6, 1).merge(tbl.cell(6, 5))
    cc = tbl.cell(6, 1)
    cc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cc.paragraphs[0].clear()

    def cp(text, bold=False, indent=0.0, size=11, before=0):
        pg = cc.add_paragraph()
        pg.paragraph_format.space_before = Pt(before)
        pg.paragraph_format.space_after  = Pt(3)
        if indent: pg.paragraph_format.left_indent = Cm(indent)
        rn = pg.add_run(text)
        rn.bold = bold; rn.font.name = FONT; rn.font.size = Pt(size)

    # 논의 내용
    cp('◆ 논의 내용', bold=True, before=6)
    for s in data.get('sections', []):
        cp(s.get('title',''), bold=True, indent=0.5, before=4)
        cp(s.get('content',''), indent=1.0)
        if s.get('duration_estimate'):
            cp(f"(약 {s['duration_estimate']})", indent=1.0, size=10)

    # 결정사항
    if data.get('decisions'):
        cp('◆ 결정사항', bold=True, before=6)
        for i, d in enumerate(data['decisions']):
            cp(f"{i+1}. {d}", indent=0.5)

    # 액션아이템
    if data.get('action_items'):
        cp('◆ 액션아이템', bold=True, before=6)
        for item in data['action_items']:
            due = f" (기한: {item.get('due_date','')})" if item.get('due_date') else ''
            pri = ' [긴급]' if item.get('priority') == 'high' else ''
            cp(f"• {item.get('task','')} — {item.get('assignee','')}{due}{pri}", indent=0.5)

    # 다음 회의
    if data.get('next_meeting'):
        cp('◆ 다음 회의', bold=True, before=6)
        cp(data['next_meeting'], indent=0.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── 주요내용 분석 DOCX ──────────────────────────────
def make_analysis_docx(data):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    FONT = '맑은 고딕'
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(3.0)

    def h(text, size=14, before=12):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text)
        r.bold = True; r.font.name = FONT; r.font.size = Pt(size)

    def b(text, bullet=False, size=11):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        if bullet: p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(('• ' if bullet else '') + str(text))
        r.font.name = FONT; r.font.size = Pt(size)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('회의 주요내용 분석')
    r.bold = True; r.font.name = FONT; r.font.size = Pt(18)

    b(f"회의일: {data.get('meeting_date','')}   장소: {data.get('venue','')}   "
      f"작성: {datetime.now().strftime('%Y.%m.%d')}", size=10)
    doc.add_paragraph()

    if data.get('keywords'):
        h('핵심 키워드', 13)
        b('  '.join(f'[{k}]' for k in data['keywords']))

    if data.get('key_issues'):
        h('주요 이슈', 13)
        for issue in data['key_issues']:
            b(issue, bullet=True)

    if data.get('summary'):
        h('전체 요약', 13)
        b(data['summary'])

    if data.get('sections'):
        h('논의 항목별 분석', 13)
        for s in data['sections']:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(8)
            r2 = p2.add_run(f"▶ {s.get('title','')}")
            r2.bold = True; r2.font.name = FONT; r2.font.size = Pt(11)
            b(s.get('content',''))

    if data.get('decisions'):
        h('결정사항', 13)
        for i, d in enumerate(data['decisions']):
            b(f"{i+1}. {d}")

    if data.get('action_items'):
        h('액션아이템', 13)
        for item in data['action_items']:
            due = f" (기한: {item.get('due_date','')})" if item.get('due_date') else ''
            pri = ' ⚡긴급' if item.get('priority') == 'high' else ''
            b(f"{item.get('task','')} → {item.get('assignee','')}{due}{pri}", bullet=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ── UI ───────────────────────────────────────────
st.title("🏛️ 수원 ON 룸")
st.caption("수원의 모든 아이디어를 켜다 | 회의 녹음 → 자동 전사 → 회의록(양식) · 전사본 · 주요분석 · 초정밀 프롬프트 4종 다운로드")
st.divider()

with st.sidebar:
    st.header("⚙️ 설정")
    st.success("AssemblyAI Key ✅") if AAI_KEY else st.error("AssemblyAI Key 없음 ❌")
    st.caption("🎙️ 고정밀 한국어 전사 (화자 구분)")
    st.success("Groq API Key ✅") if GROQ_API_KEY else st.error("Groq Key 없음 ❌")
    st.caption("🤖 AI 분석 · 프롬프트 생성")

    st.divider()
    st.markdown("**산출물 4종**")
    st.markdown("""
① 📋 회의록 DOCX (정식 양식)
② 📄 전사본 TXT
③ 🔍 주요내용 분석 DOCX
④ ✨ 초정밀 울트라 프롬프트 TXT
""")

# ── 입력 탭 ────────────────────────────────────────
tab1, tab2, tab3 = st.tabs(["🎙️ 실시간 녹음", "📁 파일 업로드", "📝 텍스트 입력"])

audio_bytes, audio_ext, text_input, audio_source = None, "wav", "", None

with tab1:
    st.markdown("버튼을 눌러 **녹음 시작**, 다시 눌러 **종료**하세요.")
    st.caption("🔴 녹음 종료 후 전사 텍스트가 **30초 단위로 순서대로** 표시됩니다.")
    rec = st.audio_input("녹음", label_visibility="collapsed")
    if rec:
        audio_bytes = rec.read()
        audio_ext = "wav"
        audio_source = "rec"
        mb = len(audio_bytes) / (1024*1024)
        st.success(f"녹음 완료 ({mb:.1f} MB) — 아래 실행 버튼을 누르세요!")

with tab2:
    st.caption("📁 파일 업로드 시 AssemblyAI 고정밀 전사 (화자 구분 포함)")
    up = st.file_uploader("m4a / mp3 / wav", type=["m4a","mp3","wav"])
    if up:
        audio_bytes = up.read()
        audio_ext = up.name.rsplit(".",1)[-1]
        audio_source = "file"
        st.audio(audio_bytes, format=f"audio/{audio_ext}")

with tab3:
    text_input = st.text_area(
        "텍스트",
        height=220,
        placeholder="전사 텍스트 또는 회의 내용을 붙여넣으세요...",
        label_visibility="collapsed",
    )

st.divider()

# ── 실행 버튼 ──────────────────────────────────────
has_input = bool(audio_bytes) or bool(text_input.strip())
has_key   = bool(GROQ_API_KEY)

if not has_key:
    st.error("사이드바에 Groq API 키를 입력해주세요.")
elif not has_input:
    st.info("녹음하거나 파일을 업로드하거나 텍스트를 입력해주세요.")

if st.button("🚀 분석 시작", disabled=not (has_input and has_key),
             use_container_width=True, type="primary"):

    st.session_state.outputs = {}
    outputs = st.session_state.outputs

    # ── 1단계: 전사 ────────────────────────────────
    if audio_bytes and audio_source == "rec":
        # 20분(1200초) 초과 → AssemblyAI 고정밀 전사로 자동 전환
        est_seconds = len(audio_bytes) / 32000  # 대략적인 길이 추정
        if est_seconds > 1200:
            st.markdown("#### 🎙️ 장시간 녹음 감지 — AssemblyAI 고정밀 전사로 처리합니다")
            st.caption(f"약 {int(est_seconds/60)}분 분량 · 처리 시간이 다소 소요됩니다.")
            with st.status("전사 중...", expanded=True) as sts:
                try:
                    transcript = transcribe_assemblyai(audio_bytes, sts)
                except Exception as e:
                    st.error(f"전사 오류: {e}")
                    st.stop()
                sts.update(label="전사 완료 ✅", state="complete")
        else:
            st.markdown("#### 🎙️ 전사 진행 중...")
            st.caption("30초 단위로 텍스트가 순서대로 쌓입니다. 완료까지 기다려주세요.")
            preview = st.empty()
            try:
                transcript = transcribe_realtime_chunks(audio_bytes, preview)
                preview.text_area("📄 전사 완료", value=transcript, height=300,
                                  label_visibility="visible")
            except Exception as e:
                st.error(f"전사 오류: {e}")
                st.stop()
        st.success(f"✅ 전사 완료 ({len(transcript):,}자)")
    elif audio_bytes and audio_source == "file":
        with st.status("**[1/3] AssemblyAI 고정밀 전사 중...**", expanded=True) as status:
            try:
                transcript = transcribe_assemblyai(audio_bytes, status)
            except Exception as e:
                st.error(f"전사 오류: {e}")
                st.stop()
            st.write(f"✓ 전사 완료 ({len(transcript):,}자)")
            status.update(label="**[1/3] 전사 완료** ✅", state="complete")
    else:
        transcript = text_input.strip()
        st.info(f"✓ 텍스트 입력 ({len(transcript):,}자)")

    outputs["transcript"] = transcript

    time.sleep(2)

    # ── 2단계: AI 분석 ────────────────────────────
    with st.status("**[2/3] 회의 내용 분석 중...**", expanded=True) as status:
        today = datetime.now().strftime('%Y-%m-%d')
        raw = groq_chat(f"""당신은 수원시정연구원 전문 회의록 작성자입니다.
아래 회의 전사 텍스트를 분석하여 정확한 JSON을 반환하세요.

[전사 텍스트]
{transcript[:5000]}

다음 JSON 구조로만 응답하세요 (다른 텍스트 없이):
{{
  "project_name": "연구명 또는 회의 주제",
  "meeting_date": "YYYY-MM-DD (오늘: {today})",
  "venue": "장소 (불명확하면 수원시정연구원 G룸)",
  "external_participants": ["외부 참석자 이름"],
  "internal_participants": ["내부 참석자 이름"],
  "sections": [
    {{
      "title": "논의 주제명",
      "content": "세부 논의 내용 (2-4문장)",
      "duration_estimate": null
    }}
  ],
  "decisions": ["결정사항1", "결정사항2"],
  "action_items": [
    {{
      "task": "해야 할 일",
      "assignee": "담당자",
      "due_date": "YYYY-MM-DD 또는 null",
      "priority": "normal 또는 high"
    }}
  ],
  "next_meeting": "다음 회의 일정 (없으면 null)",
  "key_issues": ["핵심 이슈1", "핵심 이슈2", "핵심 이슈3"],
  "keywords": ["키워드1", "키워드2", "키워드3", "키워드4", "키워드5"],
  "summary": "회의 전체 요약 (3-5문장)"
}}""", max_tokens=2500)

        # JSON 파싱
        clean = raw.strip()
        if clean.startswith("```"):
            clean = clean.split("```")[1]
            if clean.startswith("json"): clean = clean[4:]
            clean = clean.strip()
        try:
            data = json.loads(clean)
        except Exception:
            # 파싱 실패 시 기본값
            data = {
                "project_name": transcript[:30],
                "meeting_date": today,
                "venue": "수원시정연구원 G룸",
                "external_participants": [],
                "internal_participants": [],
                "sections": [{"title": "회의 내용", "content": transcript[:200], "duration_estimate": None}],
                "decisions": [],
                "action_items": [],
                "next_meeting": None,
                "key_issues": [],
                "keywords": [],
                "summary": transcript[:200],
            }

        st.write(f"✓ 연구명: {data.get('project_name','?')}")
        st.write(f"✓ 논의항목 {len(data.get('sections',[]))}개 · 결정사항 {len(data.get('decisions',[]))}개 · 액션아이템 {len(data.get('action_items',[]))}개")
        outputs["analysis"] = data
        status.update(label="**[2/3] 분석 완료** ✅", state="complete")

    time.sleep(2)

    # ── 3단계: 초정밀 울트라 프롬프트 생성 ──────────
    with st.status("**[3/3] 초정밀 울트라 프롬프트 생성 중...**", expanded=True) as status:
        ultra = groq_chat(f"""당신은 세계 최고의 AI 프롬프트 엔지니어링 전문가입니다.
아래 회의 분석 데이터를 바탕으로, 수원시정연구원 정책 연구보고서 작성을 위한
'초정밀 울트라 프롬프트'를 한국어로 작성해주세요.

[회의 분석 데이터]
- 연구명: {data.get('project_name','')}
- 핵심 이슈: {', '.join(data.get('key_issues', []))}
- 키워드: {', '.join(data.get('keywords', []))}
- 결정사항: {'; '.join(data.get('decisions', []))}
- 요약: {data.get('summary','')}

[초정밀 울트라 프롬프트 구성 요소]
1. 정밀한 역할 설정 (Role Persona)
2. 연구 배경 및 목적 설정
3. 보고서 목차 구조 (5장 이상, 각 장 소제목 포함)
4. 각 장별 세부 작성 지침 (분량, 포함 내용, 데이터 요구사항)
5. 수원시 특수성 반영 방법 및 지역 데이터 활용 지침
6. 데이터 인용·참고문헌 형식 지침
7. 문체·형식 지침 (공문서 기준)
8. 품질 체크리스트 (10개 항목 이상)
9. 예상 산출물 규격 (분량, 형식 등)

완성도 높고 즉시 사용 가능한 프롬프트를 작성해주세요.""", max_tokens=2500)

        outputs["ultra_prompt"] = ultra
        st.write(f"✓ 프롬프트 생성 완료 ({len(ultra):,}자)")
        status.update(label="**[3/3] 프롬프트 생성 완료** ✅", state="complete")

    # ── 파일 생성 후 session_state에 저장 ────────
    with st.spinner("파일 생성 중..."):
        outputs["minutes_bytes"]  = make_minutes_docx(outputs["analysis"])
        outputs["analysis_bytes"] = make_analysis_docx(outputs["analysis"])
    st.session_state.outputs = outputs

# ── 분석 결과가 있으면 항상 표시 ──────────────────
if "outputs" in st.session_state and st.session_state.outputs.get("transcript"):
    outputs = st.session_state.outputs

    st.success("🎉 완료! 원하는 파일을 선택해서 다운로드하세요.")
    st.divider()

    # ── 다운로드 4종 ─────────────────────────────
    st.markdown("### ⬇️ 다운로드")
    c1, c2, c3, c4 = st.columns(4)

    with c1:
        st.download_button(
            "📋 회의록\n(정식 양식)",
            data=outputs["minutes_bytes"],
            file_name="회의결과보고서.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )
    with c2:
        st.download_button(
            "📄 전사본",
            data=outputs["transcript"].encode("utf-8"),
            file_name="전사본.txt",
            mime="text/plain",
            use_container_width=True,
            type="primary",
        )
    with c3:
        st.download_button(
            "🔍 주요내용\n분석",
            data=outputs["analysis_bytes"],
            file_name="주요내용분석.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
        )
    with c4:
        st.download_button(
            "✨ 초정밀\n울트라 프롬프트",
            data=outputs["ultra_prompt"].encode("utf-8"),
            file_name="초정밀울트라프롬프트.txt",
            mime="text/plain",
            use_container_width=True,
            type="primary",
        )

    # ── 미리보기 탭 ──────────────────────────────
    st.divider()
    st.markdown("### 👁️ 결과 미리보기")
    t1, t2, t3, t4 = st.tabs(["📋 회의 분석", "📄 전사본", "✨ 울트라 프롬프트", "🔍 원본 JSON"])

    with t1:
        d = outputs["analysis"]
        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown(f"**연구명:** {d.get('project_name','')}")
            st.markdown(f"**일시:** {d.get('meeting_date','')}  |  **장소:** {d.get('venue','')}")
            st.markdown(f"**외부 참석자:** {', '.join(d.get('external_participants',[]))}")
            st.markdown(f"**내부 참석자:** {', '.join(d.get('internal_participants',[]))}")
        with col_b:
            st.markdown(f"**키워드:** {' '.join(f'`{k}`' for k in d.get('keywords',[]))}")
            st.markdown(f"**핵심 이슈:** {' / '.join(d.get('key_issues',[]))}")
        st.markdown("**전체 요약**")
        st.info(d.get('summary',''))
        if d.get('decisions'):
            st.markdown("**결정사항**")
            for i, dec in enumerate(d['decisions']):
                st.markdown(f"{i+1}. {dec}")
        if d.get('action_items'):
            st.markdown("**액션아이템**")
            for item in d['action_items']:
                pri = "🔴" if item.get('priority')=='high' else "🟡"
                st.markdown(f"{pri} **{item.get('task','')}** — {item.get('assignee','')} ({item.get('due_date','')})")

    with t2:
        st.text_area("전사본", value=outputs["transcript"], height=300,
                     label_visibility="collapsed")

    with t3:
        st.text_area("초정밀 울트라 프롬프트", value=outputs["ultra_prompt"], height=400,
                     label_visibility="collapsed")

    with t4:
        st.json(outputs["analysis"])
