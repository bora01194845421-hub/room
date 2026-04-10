"""
Agent 5: Policy Writer Agent
Ultra Prompt + 선행연구 → 보고서 초안 생성 → DOCX 출력
"""
import anthropic
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import DEFAULT_MODEL, OUTPUT_DIR
from templates.report_template import (
    create_base_document,
    add_title, add_heading1, add_heading2, add_heading3,
    add_body, add_bullet, add_reference_section,
    save_report,
)


class PolicyWriterAgent:
    def __init__(self):
        self.client = anthropic.Anthropic()

    # ──────────────────────────────────────────
    # 선행연구 텍스트 포매팅
    # ──────────────────────────────────────────
    def _format_references(self, references: list) -> str:
        if not references:
            return "관련 선행연구 없음"
        lines = []
        for i, ref in enumerate(references, 1):
            lines.append(
                f"[{i}] {ref.get('author', '저자미상')} ({ref.get('year', '')})."
                f" 『{ref.get('title', '')}』. "
                f"요약: {ref.get('summary', '')}"
            )
        return "\n".join(lines)

    # ──────────────────────────────────────────
    # Claude API 호출: 보고서 텍스트 생성
    # ──────────────────────────────────────────
    def _generate_report_text(
        self, ultra_prompt: str, references: list, context: dict
    ) -> str:
        refs_text = self._format_references(references)

        combined_prompt = (
            f"{ultra_prompt}\n\n"
            f"[참고 선행연구]\n{refs_text}\n\n"
            f"[보고서 제목]: {context.get('suggested_report_title', '정책 연구')}\n\n"
            "위 지침과 선행연구를 바탕으로 보고서 초안을 작성하세요. "
            "각 장(章)과 절(節)을 명확히 구분하고, 한국 행정문서 개조식 형식(□ ○ - ⦁)을 적용하세요."
        )

        message = self.client.messages.create(
            model=DEFAULT_MODEL,
            max_tokens=8000,
            messages=[{
                "role": "user",
                "content": combined_prompt
            }]
        )
        return message.content[0].text.strip()

    # ──────────────────────────────────────────
    # 텍스트 → DOCX 변환
    # ──────────────────────────────────────────
    def _generate_docx(
        self, report_text: str, context: dict, references: list
    ) -> str:
        doc = create_base_document()

        # 제목
        title = context.get("suggested_report_title", "정책 연구 보고서")
        add_title(doc, title)

        # 작성 기관/날짜 (부제)
        from datetime import datetime
        sub = doc.add_paragraph()
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        from templates.report_template import _set_run_font, SIZE_CAPTION
        r = sub.add_run(f"수원시정연구원  |  {datetime.now().strftime('%Y.%m')}")
        _set_run_font(r, SIZE_CAPTION)

        doc.add_paragraph()  # 빈 줄

        # 보고서 본문 파싱 & 삽입
        self._insert_body(doc, report_text)

        # 참고문헌
        add_reference_section(doc, references)

        return save_report(doc, "보고서초안", OUTPUT_DIR)

    # ──────────────────────────────────────────
    # 본문 파싱: 간단한 규칙 기반
    # ──────────────────────────────────────────
    def _insert_body(self, doc, text: str):
        """
        텍스트를 줄 단위로 파싱하여 스타일 적용.
        규칙:
          - "제N장" 또는 "# " → heading1
          - "## " 또는 숫자.숫자  → heading2
          - "□ " → bullet level 1
          - "○ " → bullet level 2
          - "- " → bullet level 3
          - "⦁ " → bullet level 4
          - 그 외 → body
        """
        import re
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            if re.match(r"^제\s*\d+\s*장", stripped) or stripped.startswith("# "):
                clean = re.sub(r"^#+\s*", "", stripped)
                add_heading1(doc, clean)
            elif stripped.startswith("## ") or re.match(r"^\d+\.\d+\s", stripped):
                clean = re.sub(r"^#+\s*", "", stripped)
                add_heading2(doc, clean)
            elif stripped.startswith("### "):
                clean = stripped[4:]
                add_heading3(doc, clean)
            elif stripped.startswith("□ "):
                add_bullet(doc, stripped[2:], level=1)
            elif stripped.startswith("○ "):
                add_bullet(doc, stripped[2:], level=2)
            elif stripped.startswith("- "):
                add_bullet(doc, stripped[2:], level=3)
            elif stripped.startswith("⦁ "):
                add_bullet(doc, stripped[2:], level=4)
            else:
                add_body(doc, stripped)

    # ──────────────────────────────────────────
    # 메인 실행
    # ──────────────────────────────────────────
    def run(
        self,
        ultra_prompt: str,
        references: list,
        context: dict,
    ) -> tuple[str, str]:
        """
        Returns:
            (report_path: str, report_text: str)
        """
        report_text = self._generate_report_text(ultra_prompt, references, context)
        report_path = self._generate_docx(report_text, context, references)
        return report_path, report_text
