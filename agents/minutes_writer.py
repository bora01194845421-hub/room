"""
Agent: Minutes Writer
전사 텍스트 + 맥락 → 공식 회의록 DOCX 생성
"""
import anthropic
import sys, os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import DEFAULT_MODEL, OUTPUT_DIR
from templates.report_template import (
    create_base_document, add_title, add_heading2, add_heading3,
    add_body, add_bullet, save_report, _set_run_font, SIZE_CAPTION,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

MINUTES_PROMPT = """당신은 수원시정연구원의 행정 전문가입니다.
다음 회의 전사 텍스트를 바탕으로 공식 회의록을 작성해주세요.

[회의 주제]
{main_topic}

[전사 텍스트]
{transcript}

아래 JSON 형식으로만 반환하세요:

{{
  "meeting_title": "회의명",
  "date": "회의 날짜 (텍스트에서 추출 또는 오늘 날짜)",
  "location": "장소 (텍스트에서 추출 또는 '수원시정연구원 G룸')",
  "attendees": ["참석자1", "참석자2"],
  "agenda": ["안건1", "안건2"],
  "discussion": [
    {{"topic": "논의 주제1", "content": ["논의 내용1", "논의 내용2"]}},
    {{"topic": "논의 주제2", "content": ["논의 내용1"]}}
  ],
  "decisions": ["결정사항1", "결정사항2"],
  "action_items": [
    {{"task": "할 일", "owner": "담당자", "due": "기한"}}
  ],
  "next_meeting": "다음 회의 일정 (없으면 빈 문자열)"
}}
"""


class MinutesWriterAgent:
    def __init__(self):
        self.client = anthropic.Anthropic()

    def _generate_content(self, transcript: str, context: dict) -> dict:
        import json
        msg = self.client.messages.create(
            model=DEFAULT_MODEL,
            max_tokens=3000,
            messages=[{"role": "user", "content": MINUTES_PROMPT.format(
                main_topic=context.get("main_topic", ""),
                transcript=transcript[:5000],
            )}]
        )
        raw = msg.content[0].text.strip()
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
            raw = raw.strip()
        try:
            return json.loads(raw)
        except Exception:
            return {
                "meeting_title": context.get("main_topic", "회의"),
                "date": datetime.now().strftime("%Y년 %m월 %d일"),
                "location": "수원시정연구원 G룸",
                "attendees": [],
                "agenda": [],
                "discussion": [],
                "decisions": [],
                "action_items": [],
                "next_meeting": "",
            }

    def _build_docx(self, data: dict) -> str:
        doc = create_base_document()

        # 제목
        add_title(doc, data.get("meeting_title", "회의록"))

        # 기본 정보 표
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = info.add_run(
            f"일  시: {data.get('date', '')}   |   "
            f"장  소: {data.get('location', '')}   |   "
            f"참석자: {', '.join(data.get('attendees', []))}"
        )
        _set_run_font(r, SIZE_CAPTION)
        doc.add_paragraph()

        # 안건
        add_heading2(doc, "□ 안건")
        for i, item in enumerate(data.get("agenda", []), 1):
            add_bullet(doc, f"{i}. {item}", level=2)

        # 논의 내용
        add_heading2(doc, "□ 주요 논의 내용")
        for block in data.get("discussion", []):
            add_heading3(doc, f"○ {block.get('topic', '')}")
            for line in block.get("content", []):
                add_bullet(doc, line, level=3)

        # 결정사항
        add_heading2(doc, "□ 결정사항")
        for item in data.get("decisions", []):
            add_bullet(doc, item, level=2)

        # 조치사항
        add_heading2(doc, "□ 향후 조치사항")
        for item in data.get("action_items", []):
            text = f"{item.get('task', '')}  (담당: {item.get('owner', '-')}, 기한: {item.get('due', '-')})"
            add_bullet(doc, text, level=2)

        # 다음 회의
        if data.get("next_meeting"):
            add_heading2(doc, "□ 다음 회의")
            add_body(doc, data["next_meeting"])

        return save_report(doc, "회의록", OUTPUT_DIR)

    def run(self, transcript: str, context: dict) -> tuple[str, bytes]:
        """
        Returns:
            (minutes_path, minutes_bytes)
        """
        data = self._generate_content(transcript, context)
        path = self._build_docx(data)
        with open(path, "rb") as f:
            return path, f.read()
