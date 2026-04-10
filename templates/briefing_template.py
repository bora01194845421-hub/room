"""
수원시정연구원 CEO 브리핑 DOCX 스타일 정의
A4 1페이지 분량
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

from templates.report_template import (
    FONT_MAIN, SIZE_H1, SIZE_H2, SIZE_BODY, SIZE_CAPTION,
    MARGIN_TOP, MARGIN_BOTTOM, MARGIN_LEFT, MARGIN_RIGHT,
    _set_run_font, _set_paragraph_spacing,
    BULLET_L1, BULLET_L2, BULLET_L3,
)

COLOR_HEADER = RGBColor(0x1F, 0x49, 0x7D)   # 수원시 대표 청색
COLOR_ACCENT = RGBColor(0xC0, 0x00, 0x00)   # 강조 적색


def create_briefing_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RIGHT
    doc.styles["Normal"].font.name = FONT_MAIN
    doc.styles["Normal"].font.size = SIZE_BODY
    return doc


def add_briefing_header(doc: Document, title: str, date: str = None):
    if date is None:
        date = datetime.now().strftime("%Y년 %m월 %d일")

    # 기관명
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = inst.add_run("수원시정연구원")
    _set_run_font(r, SIZE_CAPTION, color=COLOR_HEADER)
    _set_paragraph_spacing(inst, before=0, after=0)

    # 제목
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    _set_run_font(run, SIZE_H1, bold=True, color=COLOR_HEADER)
    _set_paragraph_spacing(para, before=6, after=4)

    # 날짜
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run(date)
    _set_run_font(dr, SIZE_CAPTION)
    _set_paragraph_spacing(date_para, before=0, after=8)

    # 구분선
    doc.add_paragraph("─" * 50)


def add_briefing_section(doc: Document, section_title: str, bullets: list):
    """섹션 제목 + bullet list 추가"""
    # 섹션 헤더
    hdr = doc.add_paragraph()
    run = hdr.add_run(f"{BULLET_L1} {section_title}")
    _set_run_font(run, SIZE_H2, bold=True, color=COLOR_HEADER)
    _set_paragraph_spacing(hdr, before=8, after=3)

    for item in bullets:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        r = para.add_run(f"{BULLET_L2} {item}")
        _set_run_font(r, SIZE_BODY)
        _set_paragraph_spacing(para, before=0, after=2)


def add_schedule_section(doc: Document, schedule_items: list):
    """향후 일정 섹션"""
    hdr = doc.add_paragraph()
    run = hdr.add_run(f"{BULLET_L1} 향후 추진 일정")
    _set_run_font(run, SIZE_H2, bold=True, color=COLOR_HEADER)
    _set_paragraph_spacing(hdr, before=8, after=3)

    for item in schedule_items:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        r = para.add_run(f"{BULLET_L3} {item}")
        _set_run_font(r, SIZE_BODY)
        _set_paragraph_spacing(para, before=0, after=2)


def save_briefing(doc: Document, output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(output_dir, f"CEO_브리핑_{ts}.docx")
    doc.save(path)
    return path
