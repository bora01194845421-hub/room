"""
수원시정연구원 보고서 DOCX 스타일 정의
A4 기준, 한국 행정문서 개조식
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

# ──────────────────────────────────────────────
# 폰트 / 사이즈 상수
# ──────────────────────────────────────────────
FONT_MAIN = "맑은 고딕"

SIZE_TITLE   = Pt(18)
SIZE_H1      = Pt(14)
SIZE_H2      = Pt(13)
SIZE_H3      = Pt(12)
SIZE_BODY    = Pt(11)
SIZE_CAPTION = Pt(10)

# 여백 (cm)
MARGIN_TOP    = Cm(2.5)
MARGIN_BOTTOM = Cm(2.5)
MARGIN_LEFT   = Cm(3.0)
MARGIN_RIGHT  = Cm(3.0)

# 개조식 기호
BULLET_L1 = "□"
BULLET_L2 = "○"
BULLET_L3 = "-"
BULLET_L4 = "⦁"


def _set_run_font(run, size, bold=False, color: RGBColor = None):
    run.font.name = FONT_MAIN
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT_MAIN)
    rPr.insert(0, rFonts)


def _set_paragraph_spacing(para, before=0, after=0, line_rule=None, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_rule and line:
        pf.line_spacing_rule = line_rule
        pf.line_spacing = line


def create_base_document() -> Document:
    """여백이 설정된 기본 Document 반환"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RIGHT
    # 기본 스타일 폰트
    doc.styles["Normal"].font.name = FONT_MAIN
    doc.styles["Normal"].font.size = SIZE_BODY
    return doc


def add_title(doc: Document, title: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(title)
    _set_run_font(run, SIZE_TITLE, bold=True)
    _set_paragraph_spacing(para, before=0, after=12)


def add_heading1(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(f"제{text}")
    _set_run_font(run, SIZE_H1, bold=True)
    _set_paragraph_spacing(para, before=12, after=6)


def add_heading2(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, SIZE_H2, bold=True)
    _set_paragraph_spacing(para, before=6, after=3)


def add_heading3(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, SIZE_H3, bold=True)
    _set_paragraph_spacing(para, before=3, after=3)


def add_body(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, SIZE_BODY)
    _set_paragraph_spacing(para, before=0, after=3)


def add_bullet(doc: Document, text: str, level: int = 1):
    """level: 1=□  2=○  3=-  4=⦁"""
    bullets = {1: (BULLET_L1, Cm(0)),
               2: (BULLET_L2, Cm(0.5)),
               3: (BULLET_L3, Cm(1.0)),
               4: (BULLET_L4, Cm(1.5))}
    symbol, indent = bullets.get(level, (BULLET_L3, Cm(0)))
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = indent
    run = para.add_run(f"{symbol} {text}")
    _set_run_font(run, SIZE_BODY)
    _set_paragraph_spacing(para, before=0, after=2)


def add_reference_section(doc: Document, references: list):
    if not references:
        return
    add_heading2(doc, "참고문헌")
    for i, ref in enumerate(references, 1):
        title  = ref.get("title", "")
        author = ref.get("author", "")
        year   = ref.get("year", "")
        para   = doc.add_paragraph()
        run    = para.add_run(f"{i}. {author} ({year}). 『{title}』. 수원시정연구원.")
        _set_run_font(run, SIZE_CAPTION)
        _set_paragraph_spacing(para, before=0, after=2)


def save_report(doc: Document, filename_prefix: str, output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(output_dir, f"{filename_prefix}_{ts}.docx")
    doc.save(path)
    return path
